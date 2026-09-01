"""Resume text extraction — pdf / docx / txt. Text-first, cheap, no OCR."""
from __future__ import annotations

import re
from pathlib import Path

MAX_CHARS = 9000


def parse_resume(path: str | Path) -> str:
    """Dispatch on extension -> plain text (whitespace-collapsed, capped)."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext == ".pdf":
        import fitz  # pymupdf
        text = ""
        with fitz.open(p) as doc:
            for page in doc:
                text += page.get_text()
    elif ext in (".docx", ".doc"):
        import docx
        d = docx.Document(str(p))
        text = "\n".join(par.text for par in d.paragraphs)
        for table in d.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
    elif ext in (".txt", ".md"):
        text = p.read_text(errors="replace")
    else:
        raise ValueError(f"unsupported resume type: {ext} (pdf/docx/txt only)")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()[:MAX_CHARS]


def extract_contacts(text: str) -> dict:
    """Cheap deterministic contact extraction — phone + email."""
    emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    compact = re.sub(r"[\s\-().]", "", text)
    phones = re.findall(r"(?:\+?61)?04\d{8}", compact)
    phones = [("+61" + p[-9:]) if p.startswith("04") and len(p) == 12 else p
              for p in phones]
    seen, out = set(), []
    for ph in phones:
        if ph not in seen:
            seen.add(ph)
            out.append(ph)
    return {"emails": sorted(set(emails))[:3], "phones": out[:3]}


# lines that never contain a person's name (skip when guessing)
NAME_SKIP_LINES = ("curriculum vitae", "curriculum vitae.", "cv", "resume",
                   "résumé", "professional summary", "job application")
# words that mark a line as prose rather than a name
NAME_BAD_WORDS = {"seeking", "looking", "professional", "highly", "motivated",
                  "results", "driven", "position", "opportunity", "summary",
                  "curriculum", "experienced", "enthusiastic", "dedicated",
                  "reliable", "hardworking", "skilled", "detail", "dynamic",
                  "individual", "team", "player", "self", "starter"}


def guess_name(text: str, filename: str = "") -> str:
    """Best-effort applicant name from a resume, for batch import.

    Tries the first plausible line of the parsed text (skipping emails,
    phones, separators and prose like 'Professional Summary'), then falls
    back to the filename stem ('Terry Jenkins_1788181307.txt' ->
    'Terry Jenkins'). Returns '' when nothing usable is found.
    """
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.lower().rstrip(".:") in NAME_SKIP_LINES:
            continue
        low = line.lower()
        words_raw = re.split(r"\s+", low)
        if any(w.strip(".,:;") in NAME_BAD_WORDS for w in words_raw):
            continue
        # strip emails + phone-like tokens + separators
        line = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", " ", line)
        line = re.sub(r"(\+?61\s*)?(\(?\d{2,4}\)?[\s\-.]*\d{3,4}[\s\-.]*\d{3,4})", " ", line)
        line = re.sub(r"[|·•,;:/\\]+", " ", line)
        words = [w for w in re.split(r"\s+", line)
                 if re.search(r"[A-Za-z]", w) and not re.search(r"\d", w)]
        cand = " ".join(words).strip(" .-")
        if 1 <= len(words) <= 4 and cand:
            return cand[:60]
        break  # only ever judge the first meaningful line
    stem = Path(filename).stem if filename else ""
    stem = re.sub(r"_\d+$", "", stem).replace("_", " ").strip()
    return stem[:60]
